"""DOCX resume parser using python-docx.

Extracts paragraphs, name, emails, skills, and optional phones/links.
"""

from __future__ import annotations

import io
import re
from typing import Any, Dict, List

import requests
from docx import Document

from .base import BaseParser


EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"\+?\d[\d\-\s\(\)]{7,}\d")
LINKEDIN_RE = re.compile(r"https?://[^\s,;\"]*linkedin\.com[^\s,;\"]*", re.IGNORECASE)
GITHUB_RE = re.compile(r"https?://[^\s,;\"]*github\.com[^\s,;\"]*", re.IGNORECASE)


def _split_skills(block: str) -> List[str]:
    parts = re.split(r"[,;•\n\r]+", block)
    return [p.strip() for p in parts if p and p.strip()]


class DOCXParser(BaseParser):
    """Parser for DOCX resumes using `python-docx` Document."""

    def get_source_name(self) -> str:
        name = self.source_path.name if self.source_path else self.source_path_str
        return f"resume_docx_{name}"

    def parse(self) -> Dict[str, Any]:
        if not self.validate_source():
            print(f"Source not valid or reachable: {self.source_path_str}")
            return {}

        try:
            if self.source_path is not None:
                doc = Document(str(self.source_path))
            else:
                resp = requests.get(self.source_path_str, timeout=15)
                resp.raise_for_status()
                doc = Document(io.BytesIO(resp.content))
        except Exception as e:
            print(f"Error parsing DOCX from {self.source_path_str}: {e}")
            return {}

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        full_text = "\n".join(paragraphs)

        name = paragraphs[0] if paragraphs else None

        emails = list({m.group(0) for m in EMAIL_RE.finditer(full_text)})
        phones = list({m.group(0).strip() for m in PHONE_RE.finditer(full_text)})

        links: List[str] = []
        links += LINKEDIN_RE.findall(full_text)
        links += GITHUB_RE.findall(full_text)
        links = list(dict.fromkeys(links))
        link_objs = []
        for url in links:
            link_type = 'linkedin' if 'linkedin.com' in url.lower() else (
                'github' if 'github.com' in url.lower() else 'other'
            )
            link_objs.append({'type': link_type, 'url': url})

        skills: List[str] = []
        for idx, ln in enumerate(paragraphs):
            if re.search(r"^skills?\b", ln, re.IGNORECASE) or re.search(r"technical skills", ln, re.IGNORECASE):
                after = ln.split(":", 1)
                block = after[1] if len(after) > 1 else ""
                j = idx + 1
                while (not block or len(block.split()) < 3) and j < min(len(paragraphs), idx + 8):
                    if re.match(r"^[A-Z][A-Za-z \-]{0,40}:$", paragraphs[j]):
                        break
                    block += "\n" + paragraphs[j]
                    j += 1

                skills = _split_skills(block)
                break

        skill_objects = [{'name': s} for s in skills]

        return {
            "full_name": name,
            "emails": emails,
            "phones": phones,
            "skills": skill_objects,
            "links": link_objs,
            "raw_text": full_text,
        }
